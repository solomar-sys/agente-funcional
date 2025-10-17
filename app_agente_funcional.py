import streamlit as st
from docx import Document
from datetime import datetime
import io
import time
import google.generativeai as genai

# ===================== Funções =====================

def ler_docx(file) -> str:
    doc = Document(file)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def extrair_blocos(texto):
    blocos = {
        "Requisitos Funcionais": [],
        "Perguntas de Refinamento": [],
        "Histórias": [],
        "Cenários de Teste": [],
        "Cenários de Aceite": []
    }

    secao_atual = None
    for linha in texto.splitlines():
        linha = linha.strip()
        if "requisitos funcionais" in linha.lower():
            secao_atual = "Requisitos Funcionais"
        elif "perguntas de refinamento" in linha.lower():
            secao_atual = "Perguntas de Refinamento"
        elif "histórias" in linha.lower():
            secao_atual = "Histórias"
        elif "cenários de teste" in linha.lower() or "cenarios de teste" in linha.lower():
            secao_atual = "Cenários de Teste"
        elif "cenários de aceite" in linha.lower() or "cenarios de aceite" in linha.lower():
            secao_atual = "Cenários de Aceite"
        elif secao_atual and linha:
            blocos[secao_atual].append(linha)

    # Limpeza dos itens, removendo espaços e hífens, mas preservando números e pontos
    for chave in blocos:
        itens_limpos = []
        for linha in blocos[chave]:
            if linha.startswith("•") or linha.startswith("\u2022"):
                itens_limpos.append(linha)
            else:
                linha_limpa = linha.lstrip("-* ").strip()
                itens_limpos.append(linha_limpa)
        blocos[chave] = [item for item in itens_limpos if item]

    return blocos

def escrever_docx_formatado(blocos):
    doc = Document()
    data_atual = datetime.now().strftime("%d/%m/%Y %H:%M")
    doc.add_paragraph(f"Data de geração: {data_atual}")
    doc.add_heading("Análise Funcional do Processo", level=1)

    prefixos = {
        "Requisitos Funcionais": "RF",
        "Perguntas de Refinamento": "PR",
        "Histórias": "US",
        "Cenários de Teste": "CT",
        "Cenários de Aceite": "CA"
    }

    for titulo, itens in blocos.items():
        doc.add_heading(titulo, level=2)
        prefixo = prefixos.get(titulo, "X")
        if itens:
            contador = 1
            for item in itens:
                # Verifica se item já tem prefixo com número correto para evitar duplicar
                if item.startswith(f"{prefixo}{contador}:"):
                    texto = item
                else:
                    texto = f"{prefixo}{contador}: {item}"
                doc.add_paragraph(texto, style='Normal')
                contador += 1
        else:
            doc.add_paragraph("[Nenhum conteúdo identificado]")

        # Duas linhas em branco após cada seção
        doc.add_paragraph("")
        doc.add_paragraph("")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ===================== Interface Streamlit =====================

st.set_page_config(page_title="Agente Funcional", layout="wide")
st.title("🧠 Agente Funcional - Gerador de Análise")

api_key = st.text_input("API Key do Gemini", type="password")
uploaded_file = st.file_uploader("Selecione o arquivo .docx", type=["docx"])

if st.button("Gerar Documento"):

    if not uploaded_file or not api_key:
        st.error("❌ Preencha todos os campos e selecione um arquivo.")
    else:
        try:
            # Barra de progresso e ampulheta
            progress_text = "⏳ Processando o arquivo, por favor aguarde..."
            progress_bar = st.progress(0, text=progress_text)

            # Simulação de progresso inicial
            for i in range(5):
                time.sleep(0.2)
                progress_bar.progress((i+1)*10, text=progress_text)

            # Configura API
            genai.configure(api_key=api_key)

            # Leitura e análise do arquivo
            conteudo = ler_docx(uploaded_file)
            prompt_modelo = f"""
Requisitos Funcionais  
RF1: Texto do requisito funcional 1  
RF2: Texto do requisito funcional 2  
RF3: Texto do requisito funcional 3  

Perguntas de Refinamento  
PR1: Texto da pergunta de refinamento 1  

Histórias  
US1: Como [persona], eu quero [ação] para [benefício].  

Cenários de Teste  
CT1: Cenário: [O que será testado] Dado [Pré-condição], Quando [Ação], Então [Resultado esperado].  
CT2: Cenário: [O que será testado] Dado [Pré-condição], Quando [Ação], Então [Resultado esperado].  

Cenários de Aceite  
CA1: Cenário de Aceite 1  
CA2: Cenário de Aceite 2

Não repita numeração dentro das linhas, apenas prefixo + número.

Texto do processo:
---
{conteudo}
---
"""
            modelo = genai.GenerativeModel(model_name="models/gemini-pro-latest")
            resposta = modelo.generate_content(prompt_modelo)

            # Simulação de progresso IA
            for i in range(5, 11):
                time.sleep(0.3)
                progress_bar.progress(i*10, text="🤖 IA gerando análise...")

            # Extrair e gerar docx
            blocos = extrair_blocos(resposta.text)
            buffer_docx = escrever_docx_formatado(blocos)

            st.success("✅ Documento gerado com sucesso!")
            st.download_button(
                label="📥 Baixar Documento",
                data=buffer_docx,
                file_name="Analise_Funcional.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except Exception as e:
            st.error(f"❌ Ocorreu um erro: {e}")

# para rodar: python -m streamlit run app_agente_funcional.py
